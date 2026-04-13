"""
Script tạo báo cáo Word cho dự án Hệ thống Quản lý Phòng Trọ và Gợi ý Giá Thuê
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_heading_color(paragraph, color_rgb):
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(*color_rgb)

def add_colored_heading(doc, text, level, color=(15, 23, 42)):
    heading = doc.add_heading(text, level=level)
    set_heading_color(heading, color)
    return heading

def add_table_styled(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)
        # Set header background
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '0F172A')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)

    # Data rows
    for ri, row_data in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            row_cells[ci].text = str(val)
            for para in row_cells[ci].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
        # Alternate row color
        if ri % 2 == 0:
            for cell in row_cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'EFF6FF')
                shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)

    return table


def create_report():
    doc = Document()

    # ==== TRANG BÌA ====
    # Set margins
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    # Title
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("BÁO CÁO ĐỒ ÁN")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(15, 23, 42)

    sub_title = doc.add_paragraph()
    sub_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_title.add_run("HỆ THỐNG QUẢN LÝ PHÒNG TRỌ\nVÀ GỢI Ý GIÁ THUÊ")
    r2.bold = True
    r2.font.size = Pt(18)
    r2.font.color.rgb = RGBColor(29, 78, 216)

    doc.add_paragraph()
    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = subtitle2.add_run("Mô tả kiến trúc hệ thống, cấu trúc code\nvà vị trí các chức năng chính")
    r3.font.size = Pt(13)
    r3.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()
    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_r = date_p.add_run(f"Ngày: {datetime.date.today().strftime('%d/%m/%Y')}")
    date_r.font.size = Pt(12)
    date_r.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_page_break()

    # ==== 1. TỔNG QUAN ====
    add_colored_heading(doc, "1. TỔNG QUAN DỰ ÁN", 1, (15, 23, 42))

    doc.add_paragraph(
        "Dự án xây dựng hệ thống quản lý phòng trọ tích hợp trí tuệ nhân tạo (AI) để gợi ý giá thuê. "
        "Ứng dụng được phát triển bằng Python, sử dụng Streamlit làm giao diện web, SQLAlchemy làm ORM "
        "kết nối cơ sở dữ liệu MySQL (hoặc SQLite cho môi trường test). Hệ thống phục vụ hai nhóm người dùng: "
        "Admin (chủ nhà/quản lý) và User (người thuê phòng)."
    )

    doc.add_paragraph()
    add_colored_heading(doc, "1.1. Công nghệ sử dụng", 2, (30, 64, 175))
    tech_headers = ["Công nghệ", "Mô tả", "Phiên bản"]
    tech_rows = [
        ["Python", "Ngôn ngữ lập trình chính", "3.12"],
        ["Streamlit", "Framework UI Web", "Latest"],
        ["SQLAlchemy", "ORM – Kết nối CSDL", "2.x"],
        ["MySQL / SQLite", "Hệ quản trị CSDL", "8.x / tích hợp sẵn"],
        ["Google Gemini AI", "Tích hợp AI chatbot", "gemini-2.5-flash"],
        ["ReportLab", "Xuất file PDF hợp đồng & phiếu thu", "Latest"],
        ["python-dotenv", "Quản lý biến môi trường", "Latest"],
        ["passlib", "Mã hóa mật khẩu (pbkdf2_sha256)", "Latest"],
        ["pandas", "Xử lý & hiển thị bảng dữ liệu", "Latest"],
        ["qrcode", "Tạo mã QR thanh toán MoMo", "Latest"],
    ]
    add_table_styled(doc, tech_headers, tech_rows)

    doc.add_page_break()

    # ==== 2. CẤU TRÚC THƯ MỤC ====
    add_colored_heading(doc, "2. CẤU TRÚC THƯ MỤC DỰ ÁN", 1, (15, 23, 42))

    doc.add_paragraph(
        "Toàn bộ mã nguồn được tổ chức trong thư mục tro_gia_project/tro_gia_project/ với cấu trúc như sau:"
    )

    struct_headers = ["File / Thư mục", "Loại", "Mô tả"]
    struct_rows = [
        ["app.py", "File Python (2.958 dòng)", "File chính – chứa toàn bộ logic UI, xử lý nghiệp vụ, routing"],
        ["models.py", "File Python (188 dòng)", "Định nghĩa các model ORM (dùng khi chạy độc lập)"],
        ["db.py", "File Python", "Khởi tạo engine SQLAlchemy và Base class"],
        ["auth.py", "File Python", "Hàm xác thực bổ sung (nếu dùng)"],
        ["services/billing_service.py", "Service (140 dòng)", "Tính tiền hóa đơn, tạo hóa đơn, xuất PDF phiếu thu"],
        ["services/pricing_service.py", "Service (73 dòng)", "Thuật toán gợi ý giá thuê theo khu vực & tiện ích"],
        ["services/audit_service.py", "Service", "Ghi audit log thao tác hệ thống"],
        ["static/images/", "Thư mục", "Lưu ảnh phòng upload và QR code thanh toán"],
        [".env", "Config", "Biến môi trường (DB URL, Admin credentials, Gemini API Key)"],
        ["requirements.txt", "Config", "Danh sách thư viện Python cần cài"],
        ["test.db", "SQLite DB", "CSDL SQLite dùng cho môi trường test/dev"],
    ]
    add_table_styled(doc, struct_headers, struct_rows)

    doc.add_page_break()

    # ==== 3. MÔ HÌNH DỮ LIỆU ====
    add_colored_heading(doc, "3. MÔ HÌNH DỮ LIỆU (DATABASE MODELS)", 1, (15, 23, 42))
    doc.add_paragraph(
        "Các bảng (table) được định nghĩa hai lần: một lần trong models.py (sử dụng làm tham chiếu độc lập) "
        "và một lần inline trong app.py (được dùng trực tiếp khi chạy). Cả hai có cấu trúc tương đương nhau."
    )
    doc.add_paragraph()

    model_headers = ["Bảng (Class)", "File định nghĩa", "Dòng code (app.py)", "Các trường chính"]
    model_rows = [
        ["users (User)", "app.py + models.py", "L95–L113", "user_id, full_name, username, email, phone, password_hash, role, status"],
        ["rooms (Room)", "app.py + models.py", "L116–L141", "room_id, owner_id, room_code, area_m2, khu_vuc, tang, address, current_rent, status, has_aircon, has_fridge, has_water_heater, has_balcony, has_elevator"],
        ["room_images (RoomImage)", "app.py + models.py", "L144–L153", "image_id, room_id, image_url, is_primary"],
        ["tenants (Tenant)", "app.py + models.py", "L156–L172", "tenant_id, user_id, full_name, phone, email, id_number, address"],
        ["contracts (Contract)", "app.py + models.py", "L175–L196", "contract_id, room_id, tenant_id, start_date, end_date, rent_price, deposit, payment_cycle, status, terms, digital_signature"],
        ["payments (Payment)", "app.py + models.py", "L199–L222", "payment_id, contract_id, period, amount, electricity_old/new, water_old/new, unit_price, service_fee, paid_date, method, status, note"],
        ["price_suggestions (PriceSuggestion)", "app.py + models.py", "L225–L239", "suggestion_id, room_id, suggested_price, based_on_count, algo_version, score_breakdown (JSON)"],
        ["audit_logs (AuditLog)", "app.py + models.py", "L242–L254", "audit_id, actor_user_id, entity_name, entity_id, action, old_data (JSON), new_data (JSON), changed_at"],
    ]
    add_table_styled(doc, model_headers, model_rows)

    doc.add_paragraph()
    note_p = doc.add_paragraph()
    note_r = note_p.add_run(
        "Ràng buộc CHECK quan trọng:\n"
        "• users.role: 'admin' | 'user'\n"
        "• rooms.status: 'available' | 'occupied'\n"
        "• contracts.status: 'active' | 'ended'\n"
        "• payments.status: 'paid' | 'unpaid' | 'overdue' | 'pending_verification'\n"
        "• audit_logs.action: 'insert' | 'update' | 'delete' | 'payment'"
    )
    note_r.font.size = Pt(9)
    note_r.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_page_break()

    # ==== 4. CÁC CHỨC NĂNG CHÍNH ====
    add_colored_heading(doc, "4. CÁC CHỨC NĂNG CHÍNH VÀ VỊ TRÍ CODE", 1, (15, 23, 42))
    doc.add_paragraph(
        "Dưới đây là bảng tổng hợp tất cả các chức năng của hệ thống, kèm vị trí hàm (function) "
        "trong file app.py và mô tả chi tiết."
    )

    # ---- 4.1 Hàm tiện ích & khởi động ----
    add_colored_heading(doc, "4.1. Hàm tiện ích & khởi động hệ thống", 2, (30, 64, 175))
    util_headers = ["Tên hàm", "Dòng code", "Mô tả"]
    util_rows = [
        ["serialize_model()", "L57–L69 & L572–L582", "Chuyển đổi ORM object thành dict (JSON-serializable)"],
        ["get_db()", "L508–L527", "Context manager quản lý database session (commit/rollback)"],
        ["init_database()", "L992–L1005", "Khởi tạo bảng DB và tạo admin mặc định khi lần đầu chạy"],
        ["inject_global_styles()", "L1008–L1033", "Inject CSS tùy chỉnh giao diện Streamlit"],
        ["main()", "L2906–L2957", "Hàm entry point – routing chính giữa các màn hình"],
        ["sidebar_menu()", "L1048–L1074", "Render thanh menu bên trái theo role người dùng"],
        ["hero() / section_header()", "L1036–L1045", "Render tiêu đề và banner giao diện"],
        ["money()", "L567–L569", "Format số tiền sang dạng 'X,XXX,XXX VNĐ'"],
        ["to_decimal()", "L563–L564", "Chuyển đổi giá trị sang Decimal"],
        ["room_code_text()", "L290–L294", "Lấy mã phòng dạng chuỗi an toàn"],
        ["room_label() / tenant_label()", "L297–L308", "Tạo nhãn hiển thị cho phòng và người thuê"],
        ["display_status() / db_status()", "L311–L319", "Chuyển đổi trạng thái giữa tiếng Việt và DB value"],
        ["clean_text()", "L321–L324", "Làm sạch chuỗi nhập vào (strip whitespace)"],
        ["normalize_room_search()", "L327–L338", "Chuẩn hóa phòng để tìm kiếm full-text"],
    ]
    add_table_styled(doc, util_headers, util_rows)

    doc.add_paragraph()

    # ---- 4.2 Xác thực & Phân quyền ----
    add_colored_heading(doc, "4.2. Xác thực & Phân quyền người dùng", 2, (30, 64, 175))
    auth_rows = [
        ["render_auth_screen()", "L1077–L1115", "Màn hình đăng nhập & đăng ký (2 tab)"],
        ["authenticate()", "L615–L649", "Xác thực đăng nhập (username / email / SĐT + mật khẩu)"],
        ["register_user()", "L651–L676", "Đăng ký tài khoản mới với validate đầu vào"],
        ["hash_password()", "L530–L531", "Mã hóa mật khẩu bằng pbkdf2_sha256"],
        ["verify_password()", "L534–L542", "Kiểm tra mật khẩu + hỗ trợ nâng cấp hash cũ"],
        ["current_user()", "L545–L547", "Lấy thông tin người dùng hiện tại từ session"],
        ["set_current_user()", "L550–L556", "Lưu thông tin đăng nhập vào session state"],
        ["logout()", "L559–L560", "Xóa session đăng nhập"],
        ["validate_phone()", "L585–L588", "Validate số điện thoại (0XXXXXXXXX)"],
        ["validate_email()", "L591–L594", "Validate định dạng email"],
        ["validate_password_strength()", "L597–L606", "Validate độ mạnh mật khẩu (min 8 ký tự, HOA/thường/số)"],
    ]
    add_table_styled(doc, util_headers, auth_rows)

    doc.add_paragraph()

    # ---- 4.3 Quản lý phòng ----
    add_colored_heading(doc, "4.3. Quản lý phòng (Admin)", 2, (30, 64, 175))
    room_rows = [
        ["render_rooms()", "L1211–L1457", "Màn hình quản lý phòng: 4 tab (Thêm / Sửa / Xóa / Danh sách)"],
        ["Tab 'Thêm phòng'", "L1221–L1283", "Form thêm phòng mới: mã, diện tích, khu vực, tầng, tiện ích, upload ảnh"],
        ["Tab 'Sửa phòng'", "L1285–L1382", "Hiển thị ảnh hiện tại + form sửa + upload thêm ảnh mới + xóa ảnh cũ"],
        ["Tab 'Xóa phòng'", "L1384–L1404", "Xóa phòng (chặn nếu còn hợp đồng active)"],
        ["Tab 'Danh sách'", "L1406–L1457", "Grid card hiển thị phòng với bộ lọc: từ khóa / trạng thái / khu vực"],
        ["render_user_room_catalog()", "L2165–L2287", "Màn catalog phòng cho User: grid view + detail view với ảnh, giá AI"],
    ]
    add_table_styled(doc, util_headers, room_rows)

    doc.add_paragraph()

    # ---- 4.4 Quản lý người thuê ----
    add_colored_heading(doc, "4.4. Quản lý người thuê (Admin)", 2, (30, 64, 175))
    tenant_rows = [
        ["render_tenants()", "L1461–L1632", "Màn hình quản lý người thuê: 4 tab (Thêm / Sửa / Xóa / Danh sách)"],
        ["Tab 'Thêm người thuê'", "L1481–L1518", "Form thêm: họ tên, SĐT, email, CCCD, địa chỉ, liên kết tài khoản user"],
        ["Tab 'Sửa người thuê'", "L1520–L1573", "Form sửa thông tin người thuê và liên kết lại tài khoản"],
        ["Tab 'Xóa người thuê'", "L1575–L1595", "Xóa người thuê (chặn nếu đã có hợp đồng)"],
        ["Tab 'Danh sách'", "L1597–L1632", "Bảng danh sách với bộ lọc tên/SĐT và trạng thái liên kết"],
        ["resolve_tenant_for_user()", "L455–L497", "Tự động tìm và liên kết Tenant với User theo email/SĐT"],
    ]
    add_table_styled(doc, util_headers, tenant_rows)

    doc.add_paragraph()

    # ---- 4.5 Quản lý hợp đồng ----
    add_colored_heading(doc, "4.5. Quản lý hợp đồng (Admin)", 2, (30, 64, 175))
    contract_rows = [
        ["render_contracts()", "L1636–L1857", "Màn hình quản lý hợp đồng: 3 tab"],
        ["Tab 'Tạo hợp đồng'", "L1663–L1781", "Tạo hợp đồng mới: chọn phòng trống, người thuê, ngày, giá, điều khoản"],
        ["Tab 'Kết thúc HĐ'", "L1783–L1805", "Kết thúc hợp đồng active (đổi trạng thái phòng về available)"],
        ["Tab 'Danh sách'", "L1807–L1857", "Bảng danh sách + xuất PDF hợp đồng"],
        ["validate_contract_dates()", "L609–L612", "Kiểm tra ngày kết thúc phải ≥ ngày bắt đầu"],
        ["build_contract_pdf()", "L887–L989", "Tạo file PDF hợp đồng đầy đủ (Bên A, Bên B, điều khoản)"],
        ["render_user_contracts()", "L2352–L2385", "Màn xem hợp đồng cho User – hiển thị + tải PDF"],
    ]
    add_table_styled(doc, util_headers, contract_rows)

    doc.add_paragraph()

    # ---- 4.6 Quản lý hóa đơn & thanh toán ----
    add_colored_heading(doc, "4.6. Quản lý hóa đơn & thanh toán", 2, (30, 64, 175))
    payment_rows = [
        ["render_payments()", "L1860–L1992", "Màn hình quản lý hóa đơn (Admin): 4 tab"],
        ["Tab 'Lập hóa đơn'", "L1877–L1917", "Form nhập chỉ số điện/nước, tính tổng tiền và tạo hóa đơn"],
        ["Tab 'Ghi nhận TT'", "L1919–L1952", "Xác nhận hóa đơn pending_verification + ghi nhận TT thủ công"],
        ["Tab 'Xuất phiếu thu'", "L1954–L1969", "Xuất PDF phiếu thu cho hóa đơn đã paid"],
        ["Tab 'Danh sách'", "L1971–L1992", "Bảng danh sách tất cả hóa đơn"],
        ["create_or_update_payment()", "L764–L814", "Tạo hóa đơn kỳ mới cho hợp đồng (tính tiền điện/nước/DV)"],
        ["mark_payment_paid()", "L817–L825", "Cập nhật hóa đơn thành 'paid' (ghi ngày + phương thức)"],
        ["mark_payment_pending()", "L827–L834", "Cập nhật hóa đơn thành 'pending_verification' (chờ xác nhận)"],
        ["build_receipt_pdf()", "L853–L884", "Tạo PDF phiếu thu tiền phòng"],
        ["render_user_payments()", "L2410–L2488", "Màn hóa đơn User: hiển thị unpaid/pending/paid theo nhóm"],
        ["render_user_payment_page()", "L2490–L2561", "Trang thanh toán: chọn QR/CK hoặc tiền mặt"],
        ["generate_momo_qr()", "L2388–L2407", "Tạo ảnh QR MoMo để người thuê quét thanh toán"],
    ]
    add_table_styled(doc, util_headers, payment_rows)

    doc.add_paragraph()

    # ---- 4.7 Gợi ý giá thuê (AI Pricing) ----
    add_colored_heading(doc, "4.7. Gợi ý giá thuê (AI Pricing)", 2, (30, 64, 175))
    pricing_rows = [
        ["render_price_suggestion()", "L1995–L2057", "Màn gợi ý giá thuê Admin: tính giá + lưu lịch sử"],
        ["render_user_price_suggestion()", "L2290–L2349", "Màn gợi ý giá thuê User: xem phân tích AI nội bộ"],
        ["calculate_price_for_room()", "L700–L747 (app.py)", "Thuật toán tính giá: giá TB khu vực + điều chỉnh diện tích + tiện ích + tầng"],
        ["persist_price_suggestion()", "L750–L761 (app.py)", "Lưu kết quả gợi ý giá vào bảng price_suggestions"],
        ["generate_ai_price_advice()", "L354–L383", "Tạo văn bản phân tích giá dạng ngôn ngữ tự nhiên"],
        ["pricing_service.py – calculate_price_for_room()", "services/pricing_service.py L32–L56", "Thuật toán v2: area_m2 × 180,000 × region_factor + amenity + floor_bonus"],
        ["pricing_service.py – persist_price_suggestion()", "services/pricing_service.py L59–L72", "Lưu gợi ý giá kèm số phòng tham chiếu cùng khu vực"],
    ]
    add_table_styled(doc, util_headers, pricing_rows)

    pricing_note = doc.add_paragraph()
    pricing_note.add_run(
        "\nHằng số khu vực (REGION_BENCHMARKS – app.py L260-L264):\n"
        "• Trung tâm: 3,200,000 – 5,500,000 VNĐ\n"
        "• Cận trung tâm: 2,600,000 – 4,400,000 VNĐ\n"
        "• Ngoại thành: 2,000,000 – 3,700,000 VNĐ\n\n"
        "Điều chỉnh tiện ích (AMENITY_LABELS – app.py L265-L271):\n"
        "Máy lạnh: +250,000 | Tủ lạnh: +150,000 | Bình nóng lạnh: +180,000 | Ban công: +120,000 | Thang máy: +180,000"
    ).font.size = Pt(9)

    doc.add_paragraph()

    # ---- 4.8 Audit Log ----
    add_colored_heading(doc, "4.8. Audit Log – Lịch sử thao tác hệ thống", 2, (30, 64, 175))
    audit_rows = [
        ["render_audit_logs()", "L2061–L2162", "Màn Audit Log: bảng lọc theo bảng/hành động + xem chi tiết + xuất CSV"],
        ["write_audit_log()", "L679–L697 (app.py)\nservices/audit_service.py", "Ghi một bản ghi audit log vào DB (old_data/new_data dạng JSON)"],
        ["summarize_audit_data()", "L386–L452", "Tóm tắt nội dung audit thành chuỗi ngắn để hiển thị"],
        ["audit_kv_frame()", "L341–L351", "Chuyển dict audit data thành DataFrame (Trường / Giá trị)"],
    ]
    add_table_styled(doc, util_headers, audit_rows)

    doc.add_paragraph()

    # ---- 4.9 Trợ lý AI ----
    add_colored_heading(doc, "4.9. Trợ lý AI (Tích hợp Google Gemini)", 2, (30, 64, 175))
    ai_rows = [
        ["render_ai_agent()", "L2741–L2904", "Màn Trợ lý AI Admin: chat, tác vụ tự động, báo cáo nhanh"],
        ["render_user_ai_assistant()", "L2604–L2737", "Màn Trợ lý AI User: chat hỏi về HĐ/hóa đơn/phòng trống"],
        ["Tab 'Chat AI Agent'", "L2765–L2846", "Chat với Gemini API – context thực từ DB (phòng, HĐ, công nợ)"],
        ["Tab 'Tác vụ Tự động'", "L2848–L2887", "Phân tích rủi ro công nợ, gợi ý tối ưu giá, HĐ sắp hết hạn"],
        ["Tab 'Báo cáo nhanh'", "L2889–L2903", "Biểu đồ tình trạng phòng và tài chính tổng hợp"],
        ["System prompt Admin AI", "L2779–L2793", "Cung cấp context: phòng trống, công nợ, doanh thu cho Gemini"],
        ["System prompt User AI", "L2661–L2678", "Cung cấp context: HĐ, hóa đơn nợ, phòng trống, so sánh tiện ích"],
    ]
    add_table_styled(doc, util_headers, ai_rows)

    doc.add_paragraph()

    # ---- 4.10 Quản lý User ----
    add_colored_heading(doc, "4.10. Quản lý tài khoản người dùng (Admin)", 2, (30, 64, 175))
    user_mgmt_rows = [
        ["render_user_management()", "L2564–L2602", "Màn quản lý User: xem danh sách, lọc role, thay đổi trạng thái (active/banned)"],
    ]
    add_table_styled(doc, util_headers, user_mgmt_rows)

    doc.add_paragraph()

    # ---- 4.11 Dashboard ----
    add_colored_heading(doc, "4.11. Dashboard tổng quan (Admin)", 2, (30, 64, 175))
    dash_rows = [
        ["render_dashboard()", "L1146–L1207", "Dashboard: metric (phòng, người thuê, công nợ), biểu đồ doanh thu, danh sách phòng"],
        ["build_dashboard_pdf()", "L1118–L1143", "Xuất báo cáo tổng quan hệ thống ra file PDF"],
    ]
    add_table_styled(doc, util_headers, dash_rows)

    doc.add_page_break()

    # ==== 5. PHÂN QUYỀN NGƯỜI DÙNG ====
    add_colored_heading(doc, "5. PHÂN QUYỀN NGƯỜI DÙNG", 1, (15, 23, 42))
    doc.add_paragraph(
        "Hệ thống có 2 role người dùng: admin và user. Routing được thực hiện trong hàm main() tại dòng L2920–L2953."
    )

    role_headers = ["Chức năng", "Admin", "User"]
    role_rows = [
        ["Dashboard tổng quan", "✅", "❌"],
        ["Quản lý phòng (Thêm/Sửa/Xóa)", "✅", "❌"],
        ["Xem danh sách phòng (catalog)", "✅ (full)", "✅ (xem + detail)"],
        ["Quản lý người thuê", "✅", "❌"],
        ["Quản lý hợp đồng", "✅", "Chỉ xem HĐ của mình"],
        ["Lập & xác nhận hóa đơn", "✅", "Chỉ xem + gửi yêu cầu TT"],
        ["Ghi nhận thanh toán (manual)", "✅", "❌"],
        ["Thanh toán QR/MoMo", "❌", "✅"],
        ["Gợi ý giá thuê (AI Pricing)", "✅ (full + lưu LS)", "✅ (xem + phân tích)"],
        ["Audit Log", "✅", "❌"],
        ["Quản lý tài khoản User", "✅", "❌"],
        ["Trợ lý AI (Gemini)", "✅ (Admin AI Agent)", "✅ (User AI Assistant)"],
        ["Xuất PDF hợp đồng", "✅", "✅ (HĐ của mình)"],
        ["Xuất phiếu thu PDF", "✅", "❌"],
    ]
    add_table_styled(doc, role_headers, role_rows)

    doc.add_page_break()

    # ==== 6. LUỒNG XỬ LÝ CHÍNH ====
    add_colored_heading(doc, "6. LUỒNG XỬ LÝ NGHIỆP VỤ CHÍNH", 1, (15, 23, 42))

    add_colored_heading(doc, "6.1. Luồng tạo hợp đồng", 2, (30, 64, 175))
    flows_1 = doc.add_paragraph()
    flows_1.add_run(
        "1. Admin chọn phòng trống + người thuê → nhập thông tin HĐ (ngày, giá, cọc, điều khoản)\n"
        "2. Hệ thống validate ngày bắt đầu/kết thúc → kiểm tra phòng không có HĐ active\n"
        "3. Tạo bản ghi Contract → cập nhật Room.status = 'occupied' → ghi AuditLog\n"
        "4. Tự động tìm liên kết Tenant ↔ User theo email/SĐT\n"
        "5. Hiển thị nút tải file PDF hợp đồng ngay sau khi tạo\n"
        "  → Code: render_contracts() L1690–L1781 | build_contract_pdf() L887–L989"
    )

    add_colored_heading(doc, "6.2. Luồng lập & thanh toán hóa đơn", 2, (30, 64, 175))
    flows_2 = doc.add_paragraph()
    flows_2.add_run(
        "Admin:\n"
        "1. Chọn HĐ active → nhập chỉ số điện/nước, đơn giá, phí DV → hệ thống tính tổng\n"
        "2. Tạo Payment với status='unpaid' → ghi AuditLog\n\n"
        "User:\n"
        "3. Xem hóa đơn chưa TT → nhấn 'Thanh toán'\n"
        "4. Chọn phương thức: QR/MoMo → hệ thống tạo QR code → User chuyển khoản\n"
        "5. Nhấn 'Tôi đã chuyển khoản' → status = 'pending_verification' → ghi AuditLog\n\n"
        "Admin xác nhận:\n"
        "6. Tab 'Ghi nhận TT' → mục 'Chờ xác nhận' → nhấn ✅ → status = 'paid' → ghi ngày TT\n"
        "  → Code: create_or_update_payment() L764 | mark_payment_pending() L827 | mark_payment_paid() L817"
    )

    add_colored_heading(doc, "6.3. Luồng gợi ý giá thuê", 2, (30, 64, 175))
    flows_3 = doc.add_paragraph()
    flows_3.add_run(
        "1. Chọn phòng cần tính giá\n"
        "2. Hàm calculate_price_for_room() tính:\n"
        "   • Giá TB khu vực (REGION_BENCHMARKS)\n"
        "   • Điều chỉnh diện tích = (area_m2 - 20) × 70,000\n"
        "   • Cộng tiền tiện ích (máy lạnh, tủ lạnh, BNL, ban công, thang máy)\n"
        "   • Cộng thưởng tầng = (tầng - 1) × 100,000\n"
        "   • Clamp trong [market_min, market_max]\n"
        "3. Hiển thị metric + chi tiết breakdown + lời khuyên AI (generate_ai_price_advice())\n"
        "4. Tùy chọn lưu lịch sử gợi ý vào bảng price_suggestions\n"
        "  → Code: calculate_price_for_room() L700 | render_price_suggestion() L1995"
    )

    doc.add_page_break()

    # ==== 7. CÀI ĐẶT & CHẠY ỨNG DỤNG ====
    add_colored_heading(doc, "7. CÀI ĐẶT VÀ CHẠY ỨNG DỤNG", 1, (15, 23, 42))

    add_colored_heading(doc, "7.1. Yêu cầu môi trường", 2, (30, 64, 175))
    env_rows = [
        ["APP_NAME", "Tên ứng dụng hiển thị", "Hệ thống quản lý phòng trọ..."],
        ["DATABASE_URL", "URL kết nối database", "mysql+pymysql://root:@127.0.0.1:3306/boarding_house"],
        ["ADMIN_USERNAME", "Tài khoản admin mặc định", "admin"],
        ["ADMIN_PASSWORD", "Mật khẩu admin mặc định", "Admin@123"],
        ["ADMIN_FULL_NAME", "Tên hiển thị admin", "Chủ trọ mặc định"],
        ["GEMINI_API_KEY", "API Key Google Gemini AI", "(Để sử dụng chatbot AI)"],
        ["MOMO_PHONE", "Số điện thoại MoMo nhận tiền", "0909000000"],
        ["MOMO_NAME", "Tên hiển thị MoMo", "Chủ Trọ"],
    ]
    env_headers = ["Biến môi trường", "Mô tả", "Giá trị mặc định"]
    add_table_styled(doc, env_headers, env_rows)

    doc.add_paragraph()
    add_colored_heading(doc, "7.2. Lệnh chạy ứng dụng", 2, (30, 64, 175))
    cmd_p = doc.add_paragraph()
    cmd_r = cmd_p.add_run(
        "# Cài thư viện\n"
        "pip install -r requirements.txt\n\n"
        "# Chạy ứng dụng\n"
        "cd tro_gia_project\n"
        "streamlit run app.py"
    )
    cmd_r.font.name = "Courier New"
    cmd_r.font.size = Pt(10)
    cmd_r.font.color.rgb = RGBColor(15, 118, 110)

    doc.add_page_break()

    # ==== 8. SERVICES ====
    add_colored_heading(doc, "8. SERVICES (LỚP DỊCH VỤ)", 1, (15, 23, 42))
    doc.add_paragraph(
        "Ngoài app.py, dự án tách riêng một số logic vào thư mục services/ để dễ tái sử dụng và test độc lập:"
    )

    svc_headers = ["File", "Hàm", "Mô tả"]
    svc_rows = [
        ["billing_service.py", "calculate_payment_amount()", "Tính tổng tiền hóa đơn từ chỉ số điện/nước + phí"],
        ["billing_service.py", "create_or_update_payment()", "Tạo hoặc cập nhật hóa đơn kỳ thanh toán"],
        ["billing_service.py", "mark_payment_paid()", "Đánh dấu hóa đơn đã thanh toán"],
        ["billing_service.py", "build_receipt_pdf()", "Xuất PDF phiếu thu (dùng ReportLab)"],
        ["pricing_service.py", "normalize_region()", "Chuẩn hóa tên khu vực (lowercase, map về key hợp lệ)"],
        ["pricing_service.py", "calculate_price_for_room()", "Thuật toán tính giá thuê gợi ý theo weighted scoring v1"],
        ["pricing_service.py", "persist_price_suggestion()", "Lưu kết quả gợi ý giá + đếm số phòng cùng khu vực tham chiếu"],
        ["audit_service.py", "write_audit_log()", "Ghi bản ghi audit log vào DB với thông tin entity, action, data cũ/mới"],
    ]
    add_table_styled(doc, svc_headers, svc_rows)

    doc.add_page_break()

    # ==== 9. KẾT LUẬN ====
    add_colored_heading(doc, "9. KẾT LUẬN", 1, (15, 23, 42))
    doc.add_paragraph(
        "Hệ thống Quản lý Phòng Trọ và Gợi ý Giá Thuê là một ứng dụng web đầy đủ chức năng, "
        "được xây dựng với kiến trúc monolithic rõ ràng. Toàn bộ logic nghiệp vụ chính nằm trong file "
        "app.py (~2.958 dòng code), được tổ chức theo các hàm render_*() cho từng màn hình."
    )
    doc.add_paragraph(
        "Những điểm nổi bật của dự án:"
    )
    highlights = [
        "🔒 Bảo mật: Mã hóa mật khẩu pbkdf2_sha256, kiểm tra phân quyền theo role",
        "📊 Đầy đủ CRUD: Phòng, Người thuê, Hợp đồng, Hóa đơn với validation đầu vào",
        "🤖 Tích hợp AI: Google Gemini 2.5 Flash cho cả Admin và User chatbot",
        "💡 Thuật toán giá thông minh: Tính theo khu vực thị trường + tiện ích + vị trí tầng",
        "📱 Thanh toán hiện đại: Tích hợp QR MoMo, flow pending → xác nhận bởi Admin",
        "📋 Audit Trail: Ghi log mọi thao tác thêm/sửa/xóa/thanh toán",
        "📄 Xuất tài liệu: PDF hợp đồng và phiếu thu bằng ReportLab",
        "🖼️ Upload ảnh: Quản lý ảnh phòng, lưu file local",
    ]
    for h in highlights:
        doc.add_paragraph(h, style='List Bullet')

    # Save
    output_path = r"c:\Users\Binh\Downloads\tro_gia_project\Bao_Cao_He_Thong_Quan_Ly_Phong_Tro.docx"
    doc.save(output_path)
    print(f"[OK] Da tao file bao cao Word: {output_path}")
    return output_path


if __name__ == "__main__":
    create_report()
